from memory import *
from gpt_functions import *
from helper_functions import *
from logging_functions import log_message
from search_functions import get_search_result
from wolfram_functions import get_wolframalpha_result
import datetime
import docx
import os

def generate_program(user_goal):
    role = open("data/prompts/generate_program/role.txt", "r").read()
    prompt = open("data/prompts/generate_program/prompt.txt", "r").read()

    prompt = prompt.replace("user_goal", user_goal)

    response = get_gpt_response_content(role, prompt)

    #response = get_gpt_response_content(role, prompt)
    #with open("data/prompts/generate_program/test.txt", "r") as f:
    #    response = f.readlines()

    text_lines = response.split("\n")

    program = []

    for line in text_lines:
        log_message("Parsing line: " + line + "\n")
        program.append(parse_code_line(line))
    
    return program

#Fix the program.
def fix_program(user_goal):
    role = open("data/prompts/generate_program/role.txt", "r").read()
    prompt = open("data/prompts/generate_program/prompt.txt", "r").read()

    prompt = prompt.replace("{USER_GOAL}", user_goal)

    return get_gpt_response_content(role, prompt)

#Evaluate the program.
def evaluate_program(user_goal):
    #role = open("data/prompts/generate_program/role.txt", "r").read()
    #prompt = open("data/prompts/generate_program/prompt.txt", "r").read()

    #prompt = prompt.replace("{USER_GOAL}", user_goal)

    #return get_gpt_response_content(role, prompt)
    return ""

#Declare a variable to store data in the address A-Z.
def declare(address, description, data):
    #Add data item to memory
    add_memory_item(address, description, data)

#Output data stored in memory at the address. Can output multiple address. Can output multiple times in a program.
def output_to_user(output, comparison_response = None):

    data_type = input("I have completed the goal. How would you like to output? type: terminal, docx, or txt ")

    output_path = "output/"

    if data_type == "terminal":
        #Print output to terminal
        if comparison_response != None:
            print("Comparison Response: " + comparison_response + "\n---\n")
        print(output)

    elif data_type == "docx":

        name = input("What do you want to name the file? ")

        #Save output to docx file
        date_time = datetime.datetime.now()
        date_time = date_time.strftime("%m-%d-%Y")

        if not os.path.exists(output_path + date_time):
            os.makedirs(output_path + date_time)

        doc = docx.Document()

        doc.add_heading("Placeholder Title", 0)
        doc.add_paragraph(output)

        file_name = output_path + date_time + "/" +  name + ".docx"

        doc.save(file_name)

        print("Saved to " + file_name)

        if comparison_response != None:
            doc = docx.Document()

            doc.add_heading("Placeholder Title", 0)
            doc.add_paragraph(comparison_response)

            file_name = output_path + date_time + "/" +  name + "_comparison.docx"

            doc.save(file_name)

            print("Saved to " + file_name)

        pass

    elif data_type == "txt":

        name = input("What do you want to name the file? ")

        #Save output to txt file
        date_time = datetime.datetime.now()
        date_time = date_time.strftime("%m-%d-%Y")

        if not os.path.exists(output_path + date_time):
            os.makedirs(output_path + date_time)

        file_name = output_path + date_time + "/" +  name + ".txt"

        with open(file_name, "w") as f:
            f.write(output)

        print("Saved to " + file_name)

        if comparison_response != None:
            file_name = output_path + date_time + "/" +  name + "_comparison.txt"

            with open(file_name, "w") as f:
                f.write(comparison_response)

            print("Saved to " + file_name)

        pass

    else:
        print("Invalid output type. Printing to terminal.")
        print(output)
    
#Ask for user input regarding steps or goals with “str" question. Useful if a task is vague or unclear. Useful for asking for feedback from output. Useful for creating strings on user’s personal information.
def userresponse(request, memory_text):

    if memory_text != "none":
        print("From memory:\n" + memory_text + "\n")

    return input(request + " ")

#Prompt GPT3.5, a large language model, for output with “str" description. Address arguments are optional. Useful for creating new texts. Two versions:
def queryllm(user_goal = "None provided", description = "None provided", memory_items = "None provided"):
    role = open("data/prompts/queryllm/role.txt", "r").read()
    prompt = open("data/prompts/queryllm/prompt.txt", "r").read()

    prompt = prompt.replace("request_description", description).replace("memory_items", memory_items).replace("user_goal", user_goal)

    #print("Requesting GPT with the following prompt:\n" + prompt + "\n")
    return get_gpt_response_content(role, prompt)

#Prompt GPT3.5, a large language model, for output with “str" description. Address arguments are optional. Useful for creating new texts. Two versions:
def copywriter(user_goal = "None provided", description = "None provided", memory_items = "None provided"):
    role = open("data/prompts/copywriter/role.txt", "r").read()
    prompt = open("data/prompts/copywriter/prompt.txt", "r").read()

    prompt = prompt.replace("request_description", description).replace("memory_items", memory_items).replace("user_goal", user_goal)

    #print("Requesting GPT with the following prompt:\n" + prompt + "\n")
    return get_gpt_response_content(role, prompt)

#Search the internet for text information such as facts, style guides, outlines, theory, information past 2021, etc with “str" search keywords. Address arguments are optional. Two versions:
def search(description):
    #return get_gpt_response_content(user_goal, description, memory_items)
    return get_search_result(description)

#Analyze a difficult question to something using WolframAlpha. Solve math problems like counting or calculating, science and technology related questions, society and culture related questions, and general knowledge related questions. "str" is a description with any analysis Bloom's taxonomy verb like analyze, calculate, relate, select, determine, compute, distinguish. Address arguments are optional. Two versions:
def analyze(description):
    return get_wolframalpha_result(description)

#Evaluate data to create a new text. “str" is a description with any evaluation Bloom’s taxonomy verbs like compare, contrast, revise, justify, ctritique, relate, validate, conclude, recommend, etc. Return a new evaluated string. Two versions:
def evaluate(user_goal, description, memory_items):
    #role = open("data/prompts/generate_program/role.txt", "r").read()
    #prompt = open("data/prompts/generate_program/prompt.txt", "r").read()

    return get_gpt_response_content(user_goal, description, memory_items)
